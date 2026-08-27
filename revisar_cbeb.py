#!/usr/bin/env python3
"""Revisão do manuscrito CBEB conforme parecer do revisor.

Mudanças:
1. Abstract enxugado para o limite IEEE (~250 palavras).
2. Análise comparativa ampliada: openEHR/IHE em a) + novo item d) com
   frameworks de rastreabilidade (DSCSA / EU FMD).
3. Discussão de adoção em larga escala aprofundada: mitigações concretas
   para integração com HIS, desempenho e governança.
4. Conformidade de template: 'Figure N' -> 'Fig. N', pontos finais em
   headings, keywords em itálico, 'And we' -> 'We', nbsp em caption,
   'verified: Its' -> 'verified: its', trims compensatórios.
"""
import copy
from docx import Document

SRC = "CBEB_PicSaude_FINAL_v8.docx"
DST = "CBEB_PicSaude_FINAL_v9.docx"

doc = Document(SRC)
P = doc.paragraphs

NEW_ABSTRACT = (
    "Abstract— Brazil's Unified Health System (SUS) has advanced interoperability "
    "through the National Health Data Network (RNDS) and HL7 FHIR, yet traceability "
    "remains fragmented: prescriptions, lab orders, reports, appointments, referrals, "
    "and dispensing events still lack complete audit trails, and custody often stays "
    "implicit. This paper proposes the traceable sanitary object, an architectural "
    "pattern that endows each care object with seven invariant properties: a global "
    "identifier (UUID), immutability after issuance, an append-only ledger, explicit "
    "and granular custody, a finite state machine, a canonical document with a "
    "cryptographic hash, and ownership from creation. Heterogeneous clinical objects "
    "share this core, dispensing is modeled as an act within the prescription flow, "
    "and derived objects are linked through referential lineage, mirroring Brazil's "
    "Pix payment rails. A reference implementation (Python/FastAPI, "
    "PostgreSQL/SQLite, AGPL) yields a tamper-evident audit trail and neutral public "
    "endpoints under the Brazilian General Data Protection Law (LGPD). Reproducible "
    "checks verify the pattern: 11 state machines (78 states, 99 transitions) are "
    "exhaustively checked; the canonical document's SHA-256 digest is deterministic "
    "over 1,000 recomputations and changes with any of its 13 fields; all 96 "
    "authorization scenarios pass on PostgreSQL; and none of the six public endpoints "
    "exposes a clinical field. An independent pass reproduced the reported figures. "
    "The proof of concept implements server-side ICP-Brasil signing of issued PDFs; "
    "client-side signing and an external event-publishing layer remain future work."
)

changes = []

def replace_in_runs(p, old, new, tag):
    full = "".join(r.text for r in p.runs)
    assert old in full, f"NOT FOUND [{tag}]: {old[:60]!r}"
    # simple path: replacement fits within a single run
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            changes.append(tag)
            return
    # fallback: rebuild across runs preserving first-run format
    joined = full.replace(old, new)
    p.runs[0].text = joined
    for r in p.runs[1:]:
        r.text = ""
    changes.append(tag + " (run-rebuild)")

# 1) Abstract enxugado
p = P[34]
assert p.runs[0].text.startswith("Abstract—")
p.runs[0].text = NEW_ABSTRACT
for r in p.runs[1:]:
    r.text = ""
changes.append(f"abstract trimmed 347 -> {len(NEW_ABSTRACT.split())} words")

# 2) Keywords em itálico (estilo IEEE)
p = P[35]
for r in p.runs:
    r.font.bold = False
    r.font.italic = True
changes.append("keywords -> italic")

# 3) 'And we' -> 'We'
replace_in_runs(P[41], ". And we do not merely demonstrate it;",
                ". We do not merely demonstrate it;", "fix 'And we'")

# 4) Headings com ponto final
replace_in_runs(P[46], "Electronic prescription and digital signature.",
                "Electronic prescription and digital signature", "heading period [46]")
replace_in_runs(P[136], "Ethics.", "Ethics", "heading period [136]")

# 5) three -> four paradigms (novo item d)
replace_in_runs(P[53], "three established technical paradigms",
                "four established technical paradigms", "three->four paradigms")

# 6) openEHR/IHE no item a) do posicionamento
replace_in_runs(
    P[54],
    "this architecture establishes the state machine and the custody management layer.",
    "this architecture establishes the state machine and the custody management layer. "
    "The same boundary applies to openEHR archetypes and IHE integration profiles, "
    "which standardize clinical content models and cross-enterprise document sharing "
    "but likewise leave life-cycle enforcement and custody to each implementation.",
    "openEHR/IHE comparison")

# 7) Novo item d) no posicionamento (DSCSA / EU FMD)
# Cria parágrafo novo copiando APENAS o pPr (deepcopy do parágrafo inteiro
# duplicaria commentRangeStart/End e quebraria os IDs de comentário).
from docx.oxml import OxmlElement
src_p = P[56]
new_el = OxmlElement("w:p")
if src_p._p.pPr is not None:
    new_el.append(copy.deepcopy(src_p._p.pPr))
src_p._p.addnext(new_el)
from docx.text.paragraph import Paragraph
newp = Paragraph(new_el, src_p._parent)
r1 = newp.add_run("d) ")
r1.font.italic = True
r2 = newp.add_run("Healthcare Traceability Frameworks (DSCSA / EU FMD):")
r2.font.italic = True
r3 = newp.add_run(
    " Regulatory track-and-trace regimes such as the U.S. DSCSA and the EU Falsified "
    "Medicines Directive mandate the serialization and verification of physical drug "
    "packages across the supply chain [13]-[15]. These frameworks locate a unit of "
    "product; they do not record who holds a clinical decision, nor in what state it "
    "stands. The proposed pattern operates at this clinical-object layer and is "
    "complementary to them: a dispensing event recorded under this architecture can "
    "anchor the care-level endpoint of such product-level chains."
)
changes.append("new item d) DSCSA/EU FMD comparison")

# 8) 'Figure N' -> 'Fig. N' (estilo IEEE)
replace_in_runs(P[58], "(Figure 1)", "(Fig. 1)", "Fig. 1")
replace_in_runs(P[70], "Figure 2 shows", "Fig. 2 shows", "Fig. 2")
replace_in_runs(P[70], "Figure 3 shows", "Fig. 3 shows", "Fig. 3")

# 9) Trims compensatórios (orçamento de páginas: doc está no limite de 8)
replace_in_runs(P[118],
    "—an under-count of canonical-document fields, since corrected to the 13/13 reported here",
    ", since corrected to the 13/13 reported here", "trim [118]")
replace_in_runs(P[120],
    "that an exhaustive check over the formal contract surfaces such cases is precisely",
    "surfacing such cases is precisely", "trim [120]")

# 10) nbsp no caption da Tabela III
replace_in_runs(P[131], "\xa0", "", "nbsp caption [131]")

# 11) Aprofundamento dos desafios de adoção (a, c, d)
replace_in_runs(
    P[140],
    "without demanding the complete replacement of existing hospital and pharmacy systems.",
    "without demanding the complete replacement of existing hospital and pharmacy "
    "systems. The planned mitigation is an event-publishing layer over the append-only "
    "ledger—cursor-based polling endpoints scoped by organization, per-organization "
    "webhook registration, and a canonical external event format—so that adapters "
    "consume official endpoints and never write directly to clinical tables, enabling "
    "incremental coexistence with legacy hospital information systems rather than "
    "wholesale replacement.",
    "limitations a) deepened")
replace_in_runs(
    P[142],
    "to maintain low latency during peak clinical hours.",
    "to maintain low latency during peak clinical hours. The relational design admits "
    "horizontal partitioning of the event and custody tables by time or organization, "
    "read replicas serving deterministic projections, and archival of terminally "
    "closed objects to cold storage; the absence of distributed consensus removes a "
    "major latency source by construction.",
    "limitations c) deepened")
replace_in_runs(
    P[143],
    "and a robust governance model to manage credentialing, access control, and "
    "cross-jurisdictional coordination.",
    "and a robust governance model to manage credentialing, access control, and "
    "cross-jurisdictional coordination. A federated operating model—central "
    "specification and auditing combined with regionally operated instances—could "
    "balance national consistency with local autonomy, but operator accreditation and "
    "dispute resolution across jurisdictions remain open institutional questions.",
    "limitations d) deepened")

# 12) 'verified: Its' -> 'verified: its'
replace_in_runs(P[146], "verified: Its state machines",
                "verified: its state machines", "case fix [146]")

doc.save(DST)
print("Saved:", DST)
for c in changes:
    print(" -", c)
