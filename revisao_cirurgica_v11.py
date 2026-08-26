#!/usr/bin/env python3
"""Revisão cirúrgica v10 -> v11 (modificações em vermelho).

A1. Propriedade (f): alinha com a Seção VIII (server-side PAdES-B implementado;
    futuro = assinatura client-side do documento canônico).
A2. Typos: 'relavant' -> 'relevant'; 'U.S. The DSCSA and EU' -> 'U.S. DSCSA and the EU'.
    ('It is an artifact.' já está correto na v10 — verificado.)
B1. Propriedade (d): sentença aprovada sobre unicidade de custódia como
    invariante de banco (índices únicos parciais nos dois dialetos).
C1. Âncora de snapshot das métricas: commit babc264, Aug. 24, 2026.
C2. Sentença de fecho na Conclusão (espelho da frase do produto).
D.  Table III: célula vazia de Authorization/anti-leak preenchida com o
    resultado do corpo do texto (96/96), em vermelho para revisão.
"""
import copy
from docx import Document
from docx.shared import RGBColor

SRC = "CBEB_PicSaude_FINAL_v10.docx"
DST = "CBEB_PicSaude_FINAL_v11.docx"
RED = RGBColor(0xFF, 0x00, 0x00)

doc = Document(SRC)
P = doc.paragraphs
log = []

# --- A2.1: 'relavant' -> 'relevant' [43] (runs: 'rela' + 'vant') ---
p = P[43]
assert p.runs[1].text == "rela" and p.runs[2].text == "vant"
p.runs[1].text = ""
p.runs[2].text = "relevant"
p.runs[2].font.color.rgb = RED
log.append("A2.1 relavant->relevant")

# --- A2.3: 'U.S. The DSCSA and EU' -> 'U.S. DSCSA and the EU' [58] ---
p = P[58]
assert p.runs[2].text == " The"
r3 = p.runs[3]
assert r3.text.startswith(" DSCSA and EU Falsified")
p.runs[2].text = ""
# split r3: ' DSCSA and' | ' the'(red) | ' EU Falsified...'
rest = r3.text[len(" DSCSA and"):]
r3.text = " DSCSA and"
red_el = copy.deepcopy(r3._r)
r3._r.addnext(red_el)
from docx.text.run import Run
red_run = Run(red_el, r3._parent)
red_run.text = " the"
red_run.font.color.rgb = RED
tail_el = copy.deepcopy(r3._r)
red_el.addnext(tail_el)
tail_run = Run(tail_el, r3._parent)
tail_run.text = rest
log.append("A2.3 U.S. The DSCSA -> U.S. DSCSA and the EU")

# --- B1: sentença aprovada na propriedade (d) [66] ---
p = P[66]
assert "single custodian at every moment" in p.text
r = p.add_run(
    " Uniqueness of custody is enforced at the database level by partial unique "
    "indexes in both supported dialects, rather than by application convention."
)
r.font.color.rgb = RED
log.append("B1 unicidade de custódia (invariante de banco)")

# --- A1: propriedade (f) alinhada à Seção VIII [69] ---
p = P[69]
r7 = p.runs[7]
assert r7.text == "in the current proof-of-concept, the signing step is a stub (Section VIII)."
r7.text = (
    "in the current proof-of-concept, server-side ICP-Brasil PAdES-B signing of "
    "issued PDF documents is fully implemented, while client-side signing of the "
    "canonical document remains future work (Section VIII)."
)
r7.font.color.rgb = RED
log.append("A1 propriedade (f) alinhada à VIII")

# --- C1: âncora de snapshot das métricas [103] ---
p = P[103]
assert "36,142 lines" in p.text
r = p.add_run(" All implementation figures above are as of commit babc264, Aug. 24, 2026.")
r.font.color.rgb = RED
log.append("C1 âncora as of commit babc264, Aug. 24, 2026")

# --- C2: sentença de fecho na Conclusão [147] ---
p = P[147]
assert "real care setting." in p.text
r = p.add_run(
    " In effect, the pattern supplies the substrate for the frictionless "
    "circulation of sanitary objects under citizen custody — the care-side "
    "equivalent of the rails Pix gave to payments."
)
r.font.color.rgb = RED
log.append("C2 sentença de fecho na Conclusão")

# --- D: Table III, linha Authorization/anti-leak (célula vazia) ---
t = doc.tables[2]
cell = t.rows[3].cells[1]
assert cell.text.strip() == ""
# herda tamanho de fonte da célula vizinha
from docx.shared import Pt
sib_runs = t.rows[4].cells[1].paragraphs[0].runs
sz = sib_runs[0].font.size if sib_runs else Pt(8)
cp = cell.paragraphs[0]
r = cp.add_run("96/96 scenarios pass on PostgreSQL")
r.font.color.rgb = RED
if sz: r.font.size = sz
log.append("D Table III: célula Authorization preenchida (96/96)")

doc.save(DST)
print("Saved:", DST)
for x in log:
    print(" -", x)
